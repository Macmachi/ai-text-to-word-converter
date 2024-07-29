###
# ai-text-to-word-converter
# Version : 1.0.0
# Auteur : Arnaud 
###
import sys
from PyQt6.QtWidgets import QApplication, QWidget, QTextEdit, QPushButton, QVBoxLayout, QFileDialog, QMessageBox
from PyQt6.QtGui import QFont, QColor
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class TextToWordConverter(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('ai-text-to-word-converter')
        self.setGeometry(100, 100, 800, 600)

        layout = QVBoxLayout()

        # Zone de texte
        self.textEdit = QTextEdit()
        self.textEdit.setPlaceholderText("Collez votre texte ici...")
        font = QFont("Arial", 11)
        self.textEdit.setFont(font)
        layout.addWidget(self.textEdit)

        # Bouton de génération
        self.generateButton = QPushButton('Générer le document Word')
        self.generateButton.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                border: none;
                color: white;
                padding: 15px 32px;
                text-align: center;
                font-size: 16px;
                margin: 4px 2px;
                border-radius: 8px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        self.generateButton.clicked.connect(self.generate_word_document)
        layout.addWidget(self.generateButton)

        self.setLayout(layout)

        # Définir le style de la fenêtre
        self.setStyleSheet("""
            QWidget {
                background-color: #f0f0f0;
                color: black;
            }
            QTextEdit {
                background-color: white;
                border: 2px solid #dcdcdc;
                border-radius: 8px;
                padding: 10px;
                color: black;
            }
            QMessageBox {
                color: black;
            }
        """)

    def generate_word_document(self):
        text = self.textEdit.toPlainText()
        if not text:
            QMessageBox.warning(self, "Erreur", "Veuillez entrer du texte avant de générer le document.")
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Enregistrer le document Word", "", "Documents Word (*.docx)")
        if file_path:
            if not file_path.endswith('.docx'):
                file_path += '.docx'

            # Vérifier si le fichier existe déjà
            if os.path.exists(file_path):
                reply = QMessageBox.question(self, 'Confirmer le remplacement',
                                            f"Le fichier {file_path} existe déjà. Voulez-vous le remplacer?",
                                            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                            QMessageBox.StandardButton.No)
                if reply == QMessageBox.StandardButton.No:
                    return  # L'utilisateur ne veut pas remplacer, on arrête l'opération

            try:
                self.convert_to_word(text, file_path)
                QMessageBox.information(self, "Succès", f"Document Word créé avec succès : {file_path}")
            except PermissionError:
                QMessageBox.critical(self, "Erreur", f"Permission refusée. Impossible d'écrire dans le fichier : {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Une erreur est survenue lors de la création du document : {str(e)}")

    def convert_to_word(self, text, file_path):
        doc = Document()

        # Définir les styles
        styles = doc.styles

        def get_or_add_style(style_name, style_type):
            try:
                return styles.add_style(style_name, style_type)
            except ValueError:
                return styles[style_name]

        title_style = get_or_add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True

        heading1_style = get_or_add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.size = Pt(16)
        heading1_style.font.bold = True

        heading2_style = get_or_add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True

        heading3_style = get_or_add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True

        normal_style = get_or_add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)

        list_style = get_or_add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.size = Pt(11)

        # Traiter le texte ligne par ligne
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('###'):
                p = doc.add_paragraph(line.replace('###', '').strip(), style='Heading1')
            elif line.startswith('####'):
                p = doc.add_paragraph(line.replace('####', '').strip(), style='Heading2')
            elif line.startswith('- **'):
                p = doc.add_paragraph(line.replace('- **', '').replace('**', '').strip(), style='Heading3')
            elif line.startswith('-'):
                p = doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
            elif line:
                doc.add_paragraph(line, style='Normal')

        doc.save(file_path)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = TextToWordConverter()
    ex.show()
    sys.exit(app.exec())